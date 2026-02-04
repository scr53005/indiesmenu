import qrcode
from PIL import Image, ImageDraw, ImageFont
import os
import sys
import re
from docx import Document
from docx.shared import Inches, Cm, Pt
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def _remove_table_borders(table):
    """Remove all borders from a docx table via OOXML."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    borders = OxmlElement('w:tblBorders')
    for name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        elem = OxmlElement(f'w:{name}')
        elem.set(qn('w:val'), 'nil')
        elem.set(qn('w:sz'), '0')
        elem.set(qn('w:space'), '0')
        elem.set(qn('w:color'), 'auto')
        borders.append(elem)
    tblPr.append(borders)


def _clear_cell_margins(cell):
    """Set all internal cell margins to 0 via OOXML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for name in ('top', 'left', 'bottom', 'right'):
        node = OxmlElement(f'w:{name}')
        node.set(qn('w:w'), '0')
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def _style_finders(img, modules_count, module_px):
    """
    Redraw the three finder patterns with a fully rounded outer border and a circular inner eye.
    """
    draw = ImageDraw.Draw(img)
    finder_px = 7 * module_px
    radius = int(module_px * 1.3)

    finders = [
        (0,                               0),                              # top-left
        ((modules_count - 7) * module_px,  0),                              # top-right
        (0,                               (modules_count - 7) * module_px), # bottom-left
    ]

    for fx, fy in finders:
        x1, y1 = fx, fy
        x2, y2 = fx + finder_px, fy + finder_px

        # Clear and draw the outer rounded rectangle
        draw.rectangle([x1, y1, x2, y2], fill='white')
        draw.rounded_rectangle([x1, y1, x2, y2], radius=radius, fill='black')

        # White ring: inner_radius = radius - inset keeps the border thickness uniform
        inset = module_px
        inner_radius = max(0, radius - int(inset))
        draw.rounded_rectangle(
            [x1 + inset, y1 + inset, x2 - inset, y2 - inset],
            radius=inner_radius, fill='white'
        )

        # Inner eye as circle (replaces the standard 3x3 square)
        cx, cy = fx + finder_px / 2, fy + finder_px / 2
        eye_r = module_px * 1.5  # 3-module diameter
        draw.ellipse([cx - eye_r, cy - eye_r, cx + eye_r, cy + eye_r], fill='black')


def generate_qr_codes_and_overlay(uri_file_path, tables_file_path, template_file_path, output_directory="output_qrs", test_mode=False):
    """
    Generates QR codes and overlays table numbers onto an image template,
    handling transparent backgrounds by filling them with white,
    and collects them into a single Word document with landscape pages
    in a 2x2 grid layout.

    Args:
        uri_file_path (str): Path to the text file containing the base URI.
        tables_file_path (str): Path to the CSV file containing table numbers (one per line).
        template_file_path (str): Path to the template image file (expected to be JPG).
        output_directory (str): Directory where temporary generated images will be saved before being added to the docx.
    """
    # --- 1. Read Inputs ---
    try:
        with open(uri_file_path, 'r', encoding='utf-8-sig') as f:
            base_uri = f.read().strip()
        if not base_uri:
            raise ValueError(f"'{uri_file_path}' is empty or contains only whitespace.")
    except FileNotFoundError:
        print(f"Error: URI file not found at '{uri_file_path}'. Please check the path.")
        return
    except ValueError as e:
        print(f"Error reading URI file: {e}")
        return
    except Exception as e:
        print(f"An unexpected error occurred while reading '{uri_file_path}': {e}")
        return

    try:
        with open(tables_file_path, 'r', encoding='utf-8-sig') as f:
            table_numbers = [line.strip() for line in f if line.strip()]
        if not table_numbers:
            print(f"Warning: Tables file '{tables_file_path}' is empty. No QR codes will be generated.")
            return
    except FileNotFoundError:
        print(f"Error: Tables file not found at '{tables_file_path}'. Please check the path.")
        return
    except Exception as e:
        print(f"An unexpected error occurred while reading '{tables_file_path}': {e}")
        return

    if test_mode:
        table_numbers = table_numbers[-8:]
        print(f"Test mode: limited to {len(table_numbers)} tables.")

    # Create a temporary output directory if it doesn't exist (needed for saving intermediate images for docx)
    if not os.path.exists(output_directory):
        os.makedirs(output_directory)
        print(f"Created temporary output directory: '{output_directory}'")

    # --- Define Positions and Sizes (Internal to each generated image - trial and error values) ---
    QR_CODE_SIZE_PIXELS = 470
    QR_POS_X = 508
    QR_POS_Y = 84

    TABLE_NUM_TEXT_POS_X = 760
    TABLE_NUM_TEXT_POS_Y = 17
    FONT_SIZE = 34
    TEXT_COLOR = (0, 0, 0)

    TABLE_NUM_HTTP_TEXT_POS_X = 882
    TABLE_NUM_HTTP_TEXT_POS_Y = 579
    FONT_SIZE_HTTP = 18
    TEXT_COLOR_HTTP = (0, 0, 0)

    # --- Resizing parameters for final image (before adding to DOCX) ---
    # Your template is approx. 12 cm wide, 7.17 cm high.
    # We want to place 4 per page in a 2x2 grid without resizing the individual image.
    TARGET_WIDTH_CM_FOR_DOCX = 12.0 # Fixed width for images in DOCX
    DPI = 300 # Dots Per Inch - crucial for print quality.

    # Calculate original template dimensions from the file for aspect ratio
    try:
        temp_img_for_dims = Image.open(template_file_path)
        original_template_width_pixels, original_template_height_pixels = temp_img_for_dims.size
        temp_img_for_dims.close()
    except FileNotFoundError:
        print(f"Error: Template file not found at '{template_file_path}'. Cannot determine dimensions.")
        return
    except Exception as e:
        print(f"Error opening template file '{template_file_path}' to determine dimensions: {e}")
        return

    # Calculate target height for DOCX image to maintain aspect ratio
    aspect_ratio = original_template_width_pixels / original_template_height_pixels
    TARGET_HEIGHT_CM_FOR_DOCX = TARGET_WIDTH_CM_FOR_DOCX / aspect_ratio

    print(f"Target final image size for DOCX: {TARGET_WIDTH_CM_FOR_DOCX:.2f} cm x {TARGET_HEIGHT_CM_FOR_DOCX:.2f} cm")


    # --- Font Loading ---
    font1= None
    font2= None
    try:
        font1 = ImageFont.truetype("Aptos.ttf", FONT_SIZE)
        font2 = ImageFont.truetype("Aptos.ttf", FONT_SIZE_HTTP)
        print("Using Aptos font.")
    except IOError:
        print(f"Warning: Aptos.ttf font not found. Attempting to use Arial.ttf.")
        try:
            font1 = ImageFont.truetype("arial.ttf", FONT_SIZE)
            font2 = ImageFont.truetype("arial.ttf", FONT_SIZE_HTTP)
            print("Using Arial font.")
        except IOError:
            print(f"Warning: Arial.ttf font not found. Falling back to default Pillow bitmap font.")
            font = ImageFont.load_default()
    except Exception as e:
        print(f"An unexpected error occurred while loading the font: {e}. Falling back to default Pillow bitmap font.")
        font = ImageFont.load_default()

    # --- Initialize Word Document ---
    document = Document()
    # Get the first section and set it to landscape A4
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE # Set orientation to landscape
    
    # Swap width and height for A4 landscape
    section.page_width = Cm(29.7)
    section.page_height = Cm(21)
    
    # Set margins
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(2.0) # Adjusted left/right for better layout with two images
    section.right_margin = Cm(2.0)

    generated_image_paths = [] # To store paths of images temporarily saved for DOCX

    # --- Process Each Table Number ---
    for i, table_num in enumerate(table_numbers):
        full_url = f"{base_uri}{table_num}"
        temp_output_filename = os.path.join(output_directory, f"Table_{table_num}_QR_Card_temp.jpg")

        try:
            original_template_img = Image.open(template_file_path)
            base_image = Image.new("RGB", original_template_img.size, (255, 255, 255))

            if original_template_img.mode == 'RGBA':
                base_image.paste(original_template_img, (0, 0), original_template_img)
            else:
                base_image.paste(original_template_img, (0, 0))

            current_image_to_compose = base_image

        except FileNotFoundError:
            print(f"Error: Template file not found at '{template_file_path}'. Skipping table {table_num}.")
            continue
        except Exception as e:
            print(f"Error opening or processing template file '{template_file_path}': {e}. Skipping table {table_num}.")
            continue

        # --- 2. Generate QR Code ---
        qr = qrcode.QRCode(
            version=None,
            error_correction=qrcode.constants.ERROR_CORRECT_H,
            box_size=10,
            border=0,
        )
        qr.add_data(full_url)
        qr.make(fit=True)

        qr_img_raw = qr.make_image(fill_color="black", back_color="white")
        qr_img_resized = qr_img_raw.resize((QR_CODE_SIZE_PIXELS, QR_CODE_SIZE_PIXELS), Image.LANCZOS).convert("RGB")

        # Style the finder patterns before compositing onto the template
        _style_finders(qr_img_resized, qr.modules_count, QR_CODE_SIZE_PIXELS / qr.modules_count)

        # --- 3. Image Manipulation (Pillow) ---
        current_image_to_compose.paste(qr_img_resized, (QR_POS_X, QR_POS_Y))

        draw = ImageDraw.Draw(current_image_to_compose)
        draw.text((TABLE_NUM_TEXT_POS_X, TABLE_NUM_TEXT_POS_Y), table_num, font=font1, fill=TEXT_COLOR)
        draw.text((TABLE_NUM_HTTP_TEXT_POS_X, TABLE_NUM_HTTP_TEXT_POS_Y), table_num, font=font2, fill=TEXT_COLOR_HTTP)

        # --- Prepare for DOCX insertion ---
        # The image needs to be saved at a resolution that matches the desired print size (DPI)
        # to prevent Word from resizing it unexpectedly or reducing quality.
        # Calculate target pixels based on desired CM width at DPI.
        target_width_pixels_for_docx = int(TARGET_WIDTH_CM_FOR_DOCX / 2.54 * DPI)
        target_height_pixels_for_docx = int(TARGET_HEIGHT_CM_FOR_DOCX / 2.54 * DPI) 

        # Only resize if the current image dimensions don't match the desired DPI-based pixels.
        # This ensures the temporary file has the exact dimensions expected for the print size.
        if (current_image_to_compose.width != target_width_pixels_for_docx or
            current_image_to_compose.height != target_height_pixels_for_docx):
            
            final_resized_image = current_image_to_compose.resize(
                (target_width_pixels_for_docx, target_height_pixels_for_docx),
                Image.BICUBIC # High-quality resampling
            )
        else:
            final_resized_image = current_image_to_compose # No resize needed if already at target DPI size

        # Save the image temporarily. python-docx needs a file path to insert images.
        try:
            final_resized_image.save(temp_output_filename, "JPEG", dpi=(DPI, DPI))
            print(f"Generated temporary image: {temp_output_filename}")
            generated_image_paths.append(temp_output_filename) # Add to list for docx
        except Exception as e:
            print(f"Error saving temporary image {temp_output_filename}: {e}")

    # --- Table layout dimensions (A4 landscape, derived from margins set above) ---
    AVAIL_WIDTH_CM = 29.7 - 2.0 - 2.0     # 25.7 cm
    AVAIL_HEIGHT_CM = 21.0 - 1.27 - 1.27  # 18.46 cm
    COL_WIDTH_CM = AVAIL_WIDTH_CM / 2      # 12.85 cm — wider than the image; centering within the cell creates the gap
    ROW_HEIGHT_CM = AVAIL_HEIGHT_CM / 2    # 9.23 cm  — images centered vertically within each row

    # --- Add Images to Word Document in 2x2 Grid ---
    for i in range(0, len(generated_image_paths), 4):
        table = document.add_table(rows=2, cols=2)
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _remove_table_borders(table)

        for row_idx in range(2):
            table.rows[row_idx].height = Cm(ROW_HEIGHT_CM)
            for col_idx in range(2):
                img_idx = i + row_idx * 2 + col_idx
                cell = table.cell(row_idx, col_idx)
                cell.width = Cm(COL_WIDTH_CM)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                _clear_cell_margins(cell)

                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)

                if img_idx < len(generated_image_paths):
                    run = paragraph.add_run()
                    run.add_picture(generated_image_paths[img_idx], width=Cm(TARGET_WIDTH_CM_FOR_DOCX))

    # --- Save the Word Document (retry on permission error, e.g. file open in Word) ---
    suffix = "_test" if test_mode else ""
    docx_output_filename = os.path.join(output_directory, f"Innopay_QR_Stickers_2x2_Landscape{suffix}.docx")
    save_successful = False
    while True:
        try:
            document.save(docx_output_filename)
            print(f"\nSuccessfully created Word document: {docx_output_filename}")
            save_successful = True
            break
        except PermissionError:
            print(f"\nCannot save '{docx_output_filename}' — the file is probably open in Word.")
            input("Close it in Word, then press Enter to retry...")
        except Exception as e:
            print(f"Error saving Word document: {e}")
            break

    # --- Clean up temporary image files only after a successful save ---
    if save_successful:
        for img_path in generated_image_paths:
            try:
                os.remove(img_path)
                print(f"Removed temporary file: {img_path}")
            except Exception as e:
                print(f"Error removing temporary file {img_path}: {e}")
        print(f"\nProcessed {len(table_numbers)} table numbers and compiled into '{docx_output_filename}'.")
    else:
        print(f"\nSave failed. Temporary images preserved in '{output_directory}'.")

# --- Auto-discover input files and run ---
if __name__ == "__main__":
    script_dir = os.path.dirname(os.path.abspath(__file__))
    uri_pattern = re.compile(r'^([a-z0-9]+)uri\.txt$', re.IGNORECASE)

    prefix = None
    uri_file = None
    for filename in os.listdir(script_dir):
        match = uri_pattern.match(filename)
        if match:
            prefix = match.group(1)
            uri_file = os.path.join(script_dir, filename)
            break

    if uri_file is None:
        print("Error: A file matching '<prefix>uri.txt' must be located in the script directory.")
        exit(1)

    print(f"Detected prefix: '{prefix}'")

    TABLES_FILE = os.path.join(script_dir, f"{prefix}tables.csv")
    TEMPLATE_FILE = os.path.join(script_dir, "templateQR.png")
    OUTPUT_DIR = os.path.join(script_dir, "output_qrs")

    test_mode = '--t' in sys.argv
    if test_mode:
        print("Test mode enabled (--t): generating 8 images, 2 pages.")

    generate_qr_codes_and_overlay(uri_file, TABLES_FILE, TEMPLATE_FILE, OUTPUT_DIR, test_mode=test_mode)