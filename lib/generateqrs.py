import qrcode
from PIL import Image, ImageDraw, ImageFont
import os
from docx import Document
from docx.shared import Inches, Cm # Import for precise sizing
from docx.enum.section import WD_ORIENT # Import for orientation

def generate_qr_codes_and_overlay(uri_file_path, tables_file_path, template_file_path, output_directory="output_qrs"):
    """
    Generates QR codes and overlays table numbers onto an image template,
    handling transparent backgrounds by filling them with white,
    and collects them into a single Word document with landscape pages
    in a 2x2 grid layout.

    Args:
        uri_file_path (str): Path to the text file containing the base URI.
        tables_file_path (str): Path to the CSV file containing table numbers (one per line).
        template_file_path (str): Path to the template image file (expected to be JPG).
        output_directory (str): Directory where temporary generated images will be saved before being added to the docx.
    """
    # --- 1. Read Inputs ---
    try:
        with open(uri_file_path, 'r', encoding='utf-8-sig') as f:
            base_uri = f.read().strip()
        if not base_uri:
            raise ValueError(f"'{uri_file_path}' is empty or contains only whitespace.")
    except FileNotFoundError:
        print(f"Error: URI file not found at '{uri_file_path}'. Please check the path.")
        return
    except ValueError as e:
        print(f"Error reading URI file: {e}")
        return
    except Exception as e:
        print(f"An unexpected error occurred while reading '{uri_file_path}': {e}")
        return

    try:
        with open(tables_file_path, 'r', encoding='utf-8-sig') as f:
            table_numbers = [line.strip() for line in f if line.strip()]
        if not table_numbers:
            print(f"Warning: Tables file '{tables_file_path}' is empty. No QR codes will be generated.")
            return
    except FileNotFoundError:
        print(f"Error: Tables file not found at '{tables_file_path}'. Please check the path.")
        return
    except Exception as e:
        print(f"An unexpected error occurred while reading '{tables_file_path}': {e}")
        return

    # Create a temporary output directory if it doesn't exist (needed for saving intermediate images for docx)
    if not os.path.exists(output_directory):
        os.makedirs(output_directory)
        print(f"Created temporary output directory: '{output_directory}'")

    # --- Define Positions and Sizes (Internal to each generated image - trial and error values) ---
    QR_CODE_SIZE_PIXELS = 545
    QR_POS_X = 861
    QR_POS_Y = 176

    TABLE_NUM_TEXT_POS_X = 862
    TABLE_NUM_TEXT_POS_Y = 783
    FONT_SIZE = 34
    TEXT_COLOR = (0, 0, 0)

    # --- Resizing parameters for final image (before adding to DOCX) ---
    # Your template is approx. 12 cm wide, 7.17 cm high.
    # We want to place 4 per page in a 2x2 grid without resizing the individual image.
    TARGET_WIDTH_CM_FOR_DOCX = 12.0 # Fixed width for images in DOCX
    DPI = 300 # Dots Per Inch - crucial for print quality.

    # Calculate original template dimensions from the file for aspect ratio
    try:
        temp_img_for_dims = Image.open(template_file_path)
        original_template_width_pixels, original_template_height_pixels = temp_img_for_dims.size
        temp_img_for_dims.close()
    except FileNotFoundError:
        print(f"Error: Template file not found at '{template_file_path}'. Cannot determine dimensions.")
        return
    except Exception as e:
        print(f"Error opening template file '{template_file_path}' to determine dimensions: {e}")
        return

    # Calculate target height for DOCX image to maintain aspect ratio
    aspect_ratio = original_template_width_pixels / original_template_height_pixels
    TARGET_HEIGHT_CM_FOR_DOCX = TARGET_WIDTH_CM_FOR_DOCX / aspect_ratio

    print(f"Target final image size for DOCX: {TARGET_WIDTH_CM_FOR_DOCX:.2f} cm x {TARGET_HEIGHT_CM_FOR_DOCX:.2f} cm")


    # --- Font Loading ---
    font = None
    try:
        font = ImageFont.truetype("Aptos.ttf", FONT_SIZE)
        print("Using Aptos font.")
    except IOError:
        print(f"Warning: Aptos.ttf font not found. Attempting to use Arial.ttf.")
        try:
            font = ImageFont.truetype("arial.ttf", FONT_SIZE)
            print("Using Arial font.")
        except IOError:
            print(f"Warning: Arial.ttf font not found. Falling back to default Pillow bitmap font.")
            font = ImageFont.load_default()
    except Exception as e:
        print(f"An unexpected error occurred while loading the font: {e}. Falling back to default Pillow bitmap font.")
        font = ImageFont.load_default()

    # --- Initialize Word Document ---
    document = Document()
    # Get the first section and set it to landscape A4
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE # Set orientation to landscape
    
    # Swap width and height for A4 landscape
    section.page_width = Cm(29.7)
    section.page_height = Cm(21)
    
    # Set margins
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    section.left_margin = Cm(2.0) # Adjusted left/right for better layout with two images
    section.right_margin = Cm(2.0)

    generated_image_paths = [] # To store paths of images temporarily saved for DOCX

    # --- Process Each Table Number ---
    for i, table_num in enumerate(table_numbers):
        full_url = f"{base_uri}{table_num}"
        temp_output_filename = os.path.join(output_directory, f"Table_{table_num}_QR_Card_temp.jpg")

        try:
            original_template_img = Image.open(template_file_path)
            base_image = Image.new("RGB", original_template_img.size, (255, 255, 255))

            if original_template_img.mode == 'RGBA':
                base_image.paste(original_template_img, (0, 0), original_template_img)
            else:
                base_image.paste(original_template_img, (0, 0))

            current_image_to_compose = base_image

        except FileNotFoundError:
            print(f"Error: Template file not found at '{template_file_path}'. Skipping table {table_num}.")
            continue
        except Exception as e:
            print(f"Error opening or processing template file '{template_file_path}': {e}. Skipping table {table_num}.")
            continue

        # --- 2. Generate QR Code ---
        qr = qrcode.QRCode(
            version=None,
            error_correction=qrcode.constants.ERROR_CORRECT_H,
            box_size=10,
            border=0,
        )
        qr.add_data(full_url)
        qr.make(fit=True)

        qr_img_raw = qr.make_image(fill_color="black", back_color="white")
        qr_img_resized = qr_img_raw.resize((QR_CODE_SIZE_PIXELS, QR_CODE_SIZE_PIXELS), Image.LANCZOS)

        # --- 3. Image Manipulation (Pillow) ---
        current_image_to_compose.paste(qr_img_resized, (QR_POS_X, QR_POS_Y))

        draw = ImageDraw.Draw(current_image_to_compose)
        draw.text((TABLE_NUM_TEXT_POS_X, TABLE_NUM_TEXT_POS_Y), table_num, font=font, fill=TEXT_COLOR)

        # --- Prepare for DOCX insertion ---
        # The image needs to be saved at a resolution that matches the desired print size (DPI)
        # to prevent Word from resizing it unexpectedly or reducing quality.
        # Calculate target pixels based on desired CM width at DPI.
        target_width_pixels_for_docx = int(TARGET_WIDTH_CM_FOR_DOCX / 2.54 * DPI)
        target_height_pixels_for_docx = int(TARGET_HEIGHT_CM_FOR_DOCX / 2.54 * DPI) 

        # Only resize if the current image dimensions don't match the desired DPI-based pixels.
        # This ensures the temporary file has the exact dimensions expected for the print size.
        if (current_image_to_compose.width != target_width_pixels_for_docx or
            current_image_to_compose.height != target_height_pixels_for_docx):
            
            final_resized_image = current_image_to_compose.resize(
                (target_width_pixels_for_docx, target_height_pixels_for_docx),
                Image.BICUBIC # High-quality resampling
            )
        else:
            final_resized_image = current_image_to_compose # No resize needed if already at target DPI size

        # Save the image temporarily. python-docx needs a file path to insert images.
        try:
            final_resized_image.save(temp_output_filename, "JPEG", dpi=(DPI, DPI))
            print(f"Generated temporary image: {temp_output_filename}")
            generated_image_paths.append(temp_output_filename) # Add to list for docx
        except Exception as e:
            print(f"Error saving temporary image {temp_output_filename}: {e}")

    # --- Add Images to Word Document in 2x2 Grid ---
    for i in range(0, len(generated_image_paths), 4): # Step by 4 for 4 images per page
        # Row 1
        paragraph_row1 = document.add_paragraph()
        
        # Image 1 (top-left)
        if i < len(generated_image_paths):
            run1 = paragraph_row1.add_run()
            run1.add_picture(generated_image_paths[i], width=Cm(TARGET_WIDTH_CM_FOR_DOCX))
            run1.add_text("      ") # Add horizontal space

        # Image 2 (top-right)
        if i + 1 < len(generated_image_paths):
            run2 = paragraph_row1.add_run()
            run2.add_picture(generated_image_paths[i+1], width=Cm(TARGET_WIDTH_CM_FOR_DOCX))
        
        # Add a small vertical space between rows. A new paragraph implicitly adds some.
        # For more precise vertical spacing between rows, one might use a table,
        # but for simplicity, we'll rely on paragraph spacing and ensure a new paragraph for row 2.
        
        # Row 2
        paragraph_row2 = document.add_paragraph()

        # Image 3 (bottom-left)
        if i + 2 < len(generated_image_paths):
            run3 = paragraph_row2.add_run()
            run3.add_picture(generated_image_paths[i+2], width=Cm(TARGET_WIDTH_CM_FOR_DOCX))
            run3.add_text("  ") # Add horizontal space

        # Image 4 (bottom-right)
        if i + 3 < len(generated_image_paths):
            run4 = paragraph_row2.add_run()
            run4.add_picture(generated_image_paths[i+3], width=Cm(TARGET_WIDTH_CM_FOR_DOCX))

        # Add a page break after every completed set of 4 images, unless it's the very last set.
        if (i + 4 <= len(generated_image_paths)): # Only add if there are more images to follow on a new page
            document.add_page_break()

    # --- Save the Word Document ---
    docx_output_filename = os.path.join(output_directory, "Innopay_QR_Stickers_2x2_Landscape.docx")
    try:
        document.save(docx_output_filename)
        print(f"\nSuccessfully created Word document: {docx_output_filename}")
    except Exception as e:
        print(f"Error saving Word document: {e}")

    # --- Clean up temporary image files ---
    for img_path in generated_image_paths:
        try:
            os.remove(img_path)
            print(f"Removed temporary file: {img_path}")
        except Exception as e:
            print(f"Error removing temporary file {img_path}: {e}")

    print(f"\nProcessed {len(table_numbers)} table numbers and compiled into '{docx_output_filename}'.")

# --- Example Usage ---
if __name__ == "__main__":
    URI_FILE = "indiesuri.txt"
    TABLES_FILE = "indiestables.csv"
    TEMPLATE_FILE = "TableQR-template.jpg"
    OUTPUT_DIR = "output_qrs" # This will now be a temporary directory for image processing

    generate_qr_codes_and_overlay(URI_FILE, TABLES_FILE, TEMPLATE_FILE, OUTPUT_DIR)